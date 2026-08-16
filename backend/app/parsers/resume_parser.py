import re
import io
from typing import Dict, List, Any, Tuple, Optional
from app.schemas.matcher_schemas import ResumeStructure, ContactInfo


class ResumeParser:
    """
    Robust resume parser supporting PDF, DOCX, and TXT.
    Extracts text, segments document into semantic sections, and analyzes ATS formatting hazards.
    """

    SECTION_PATTERNS = {
        "summary": [
            r"^(?:professional\s+)?summary",
            r"^(?:career\s+)?objective",
            r"^profile",
            r"^about(?:\s+me)?",
            r"^executive\s+summary",
        ],
        "skills": [
            r"^(?:technical\s+)?skills",
            r"^core\s+competencies",
            r"^technologies(?:\s+used)?",
            r"^skills\s+&\s+tools",
            r"^technical\s+proficiencies",
            r"^areas\s+of\s+expertise",
        ],
        "experience": [
            r"^(?:work\s+|professional\s+|employment\s+)?experience",
            r"^employment\s+history",
            r"^work\s+history",
            r"^relevant\s+experience",
        ],
        "projects": [
            r"^(?:key\s+|featured\s+|notable\s+|personal\s+)?projects",
            r"^academic\s+projects",
            r"^portfolio",
        ],
        "education": [
            r"^education(?:\s+and\s+qualifications)?",
            r"^academic\s+background",
            r"^degrees",
            r"^educational\s+background",
        ],
        "certifications": [
            r"^certifications?",
            r"^licenses(?:\s+and\s+certifications)?",
            r"^credentials",
            r"^courses\s+&\s+certifications",
        ],
        "achievements": [
            r"^achievements",
            r"^awards(?:\s+and\s+honors)?",
            r"^accomplishments",
            r"^key\s+accomplishments",
        ],
    }

    @classmethod
    def parse_file(cls, file_bytes: bytes, filename: str, file_type: Optional[str] = None) -> Tuple[str, ResumeStructure]:
        raw_text = ""
        ext = filename.lower().split(".")[-1] if filename else ""
        if file_type:
            file_type = file_type.lower()
        else:
            file_type = ext

        if "pdf" in file_type or ext == "pdf":
            raw_text = cls._extract_from_pdf(file_bytes)
        elif "docx" in file_type or ext == "docx" or "word" in file_type:
            raw_text = cls._extract_from_docx(file_bytes)
        elif "txt" in file_type or ext in ["txt", "text"]:
            raw_text = cls._extract_from_txt(file_bytes)
        else:
            # Fallback to plain text decode
            raw_text = cls._extract_from_txt(file_bytes)

        if not raw_text.strip():
            raise ValueError("Extracted resume text is empty or could not be read.")

        structured_data = cls.parse_raw_text(raw_text)
        return raw_text, structured_data

    @classmethod
    def _extract_from_pdf(cls, file_bytes: bytes) -> str:
        text_parts = []
        # Try pdfplumber first for better layout preservation
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text(layout=True) or page.extract_text()
                    if extracted:
                        text_parts.append(extracted)
        except Exception:
            # Fallback to pypdf
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    txt = page.extract_text()
                    if txt:
                        text_parts.append(txt)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF: {str(e)}")

        return "\n".join(text_parts).strip()

    @classmethod
    def _extract_from_docx(cls, file_bytes: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @classmethod
    def _extract_from_txt(cls, file_bytes: bytes) -> str:
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="ignore").strip()

    @classmethod
    def parse_raw_text(cls, raw_text: str) -> ResumeStructure:
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        contact_info = cls._extract_contact_info(raw_text, lines)
        sections = cls._segment_sections(raw_text)
        
        # Candidate name inference
        candidate_name = contact_info.name
        if not candidate_name and lines:
            first_line = lines[0]
            if len(first_line.split()) <= 4 and not re.search(r"[@\d+:/]", first_line):
                candidate_name = first_line

        # Parse skills
        skills_raw = sections.get("skills", "")
        skills_list = cls._extract_skills_list(skills_raw, raw_text)

        # Parse experience items
        exp_raw = sections.get("experience", "")
        work_experience = cls._extract_experience_blocks(exp_raw)

        # Parse projects
        proj_raw = sections.get("projects", "")
        projects = cls._extract_project_blocks(proj_raw)

        # Parse education
        edu_raw = sections.get("education", "")
        education = cls._extract_education_blocks(edu_raw)

        # Parse certs & achievements
        certs = [c.strip() for c in sections.get("certifications", "").splitlines() if c.strip() and len(c.strip()) > 3]
        achievements = [a.strip() for a in sections.get("achievements", "").splitlines() if a.strip() and len(a.strip()) > 3]

        # Detect ATS Hazards
        hazards = cls._detect_ats_hazards(raw_text, sections, contact_info)

        return ResumeStructure(
            candidate_name=candidate_name,
            contact_info=contact_info,
            professional_summary=sections.get("summary"),
            skills=skills_list,
            work_experience=work_experience,
            projects=projects,
            education=education,
            certifications=certs,
            achievements=achievements,
            raw_sections=sections,
            detected_hazards=hazards,
        )

    @classmethod
    def _extract_contact_info(cls, text: str, lines: List[str]) -> ContactInfo:
        info = ContactInfo()
        
        # Email
        email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b", text)
        if email_match:
            info.email = email_match.group(0)

        # Phone
        phone_match = re.search(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b", text)
        if phone_match:
            info.phone = phone_match.group(0)

        # LinkedIn
        linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)", text, re.IGNORECASE)
        if linkedin_match:
            info.linkedin = linkedin_match.group(0)

        # GitHub
        github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9_-]+)", text, re.IGNORECASE)
        if github_match:
            info.github = github_match.group(0)

        # Location heuristic
        loc_match = re.search(r"\b([A-Z][a-zA-Z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?|[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z]+)\b", text)
        if loc_match:
            info.location = loc_match.group(0)

        # Name heuristic from first 3 lines
        for line in lines[:3]:
            if "@" not in line and not re.search(r"\d", line) and len(line.split()) in [2, 3, 4]:
                if not re.search(r"(resume|curriculum|vitae|developer|engineer|manager)", line, re.IGNORECASE):
                    info.name = line.strip()
                    break

        return info

    @classmethod
    def _segment_sections(cls, text: str) -> Dict[str, str]:
        sections: Dict[str, str] = {}
        lines = text.splitlines()
        
        current_section = "header"
        section_lines: Dict[str, List[str]] = {current_section: []}

        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                continue

            # Check if this line is a section heading
            matched_sec = None
            if len(trimmed) < 45:
                clean_heading = re.sub(r"^[#*_\-–\s]+|[#*_\-–:\s]+$", "", trimmed).strip().lower()
                for sec_name, patterns in cls.SECTION_PATTERNS.items():
                    if any(re.search(pat, clean_heading, re.IGNORECASE) for pat in patterns):
                        matched_sec = sec_name
                        break

            if matched_sec:
                current_section = matched_sec
                if current_section not in section_lines:
                    section_lines[current_section] = []
            else:
                if current_section not in section_lines:
                    section_lines[current_section] = []
                section_lines[current_section].append(trimmed)

        for sec, lns in section_lines.items():
            sections[sec] = "\n".join(lns).strip()

        return sections

    @classmethod
    def _extract_skills_list(cls, skills_text: str, full_text: str) -> List[str]:
        skills = set()
        source_text = skills_text if len(skills_text) > 20 else full_text
        
        # Split by comma, bullet points, pipes, slashes, or newlines
        tokens = re.split(r"[,•|/\n\t;]+", source_text)
        for token in tokens:
            cleaned = re.sub(r"^[•\-\*\s]+|[:.]+$", "", token).strip()
            # Clean off labels like "Languages:", "Frameworks:"
            cleaned = re.sub(r"^(?:languages|frameworks|tools|databases|cloud|libraries|technologies):\s*", "", cleaned, flags=re.IGNORECASE).strip()
            if 1 < len(cleaned) < 40 and not re.search(r"(responsible|managed|developed|experience with|years)", cleaned, re.IGNORECASE):
                skills.add(cleaned)

        return sorted(list(skills))

    @classmethod
    def _extract_experience_blocks(cls, exp_text: str) -> List[Dict[str, Any]]:
        blocks = []
        lines = exp_text.splitlines()
        current_block = {"role": "", "company": "", "duration": "", "bullets": []}
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # Check for date ranges e.g. 2021 - Present, Jan 2020 - Dec 2022
            date_match = re.search(r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{4})\s*[-–to\s]+\s*(?:Present|Current|\d{4}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec))", line_str, re.IGNORECASE)
            
            if date_match and (len(line_str) < 90):
                if current_block["role"] or current_block["bullets"]:
                    blocks.append(current_block)
                current_block = {
                    "role": line_str.replace(date_match.group(0), "").strip(" ,-|"),
                    "company": "",
                    "duration": date_match.group(0).strip(),
                    "bullets": []
                }
            elif line_str.startswith(("•", "-", "*", "–")) or re.match(r"^\d+\.", line_str):
                bullet = re.sub(r"^[•\-*–\d.]+\s*", "", line_str).strip()
                if bullet:
                    current_block["bullets"].append(bullet)
            else:
                if not current_block["role"] and len(line_str) < 60:
                    current_block["role"] = line_str
                else:
                    current_block["bullets"].append(line_str)

        if current_block["role"] or current_block["bullets"]:
            blocks.append(current_block)

        return blocks

    @classmethod
    def _extract_project_blocks(cls, proj_text: str) -> List[Dict[str, Any]]:
        projects = []
        lines = proj_text.splitlines()
        current_proj = {"name": "", "description": "", "tech_stack": [], "bullets": []}

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith(("•", "-", "*", "–")):
                current_proj["bullets"].append(line_str.lstrip("•-*– "))
            else:
                if current_proj["name"]:
                    projects.append(current_proj)
                current_proj = {"name": line_str, "description": "", "tech_stack": [], "bullets": []}

        if current_proj["name"]:
            projects.append(current_proj)
        return projects

    @classmethod
    def _extract_education_blocks(cls, edu_text: str) -> List[Dict[str, Any]]:
        edu_list = []
        lines = edu_text.splitlines()
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if re.search(r"(bachelor|master|phd|b\.s|m\.s|b\.e|b\.tech|m\.tech|degree|university|college|institute|gpa)", line_str, re.IGNORECASE):
                edu_list.append({"degree_or_institution": line_str})
        return edu_list

    @classmethod
    def _detect_ats_hazards(cls, raw_text: str, sections: Dict[str, str], contact: ContactInfo) -> List[str]:
        hazards = []
        
        # Missing essential contact
        if not contact.email:
            hazards.append("Missing or non-extractable email address")
        if not contact.phone:
            hazards.append("Missing or non-extractable phone number")

        # Section presence
        if "experience" not in sections or len(sections["experience"]) < 50:
            hazards.append("Work Experience section heading not standard or missing")
        if "skills" not in sections or len(sections["skills"]) < 10:
            hazards.append("Dedicated Skills section heading not detected")
        if "education" not in sections or len(sections["education"]) < 10:
            hazards.append("Education section heading not detected")

        # Table or pipe indicators
        pipe_count = raw_text.count("|")
        if pipe_count > 15:
            hazards.append("Excessive pipe/table formatting detected which may hinder standard ATS parsers")

        # Suspiciously short or non-standard characters
        if len(raw_text.split()) < 100:
            hazards.append("Resume length is unusually short (less than 100 words)")
        elif len(raw_text.split()) > 1600:
            hazards.append("Resume exceeds recommended length for early screening (over 1600 words)")

        # Date presence
        has_dates = bool(re.search(r"\b(20\d{2}|19\d{2})\b", raw_text))
        if not has_dates:
            hazards.append("No standard year dates found in work history")

        return hazards
